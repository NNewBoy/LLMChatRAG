"""文档解析模块 - 支持 PDF/Word/HTML/TXT 格式"""

import os
from pathlib import Path
from utils.logger import logger


class DocumentParser:
    """文档解析器，支持 PDF、Word、HTML、TXT 格式"""

    SUPPORTED_TYPES = {"pdf", "word", "html", "txt"}

    # 文件扩展名到类型的映射
    EXTENSION_MAP = {
        ".pdf": "pdf",
        ".docx": "word",
        ".doc": "word",
        ".html": "html",
        ".htm": "html",
        ".txt": "txt",
        ".md": "txt",
    }

    def get_file_type(self, filename: str) -> str | None:
        """根据文件扩展名获取文件类型"""
        ext = Path(filename).suffix.lower()
        return self.EXTENSION_MAP.get(ext)

    def parse(self, file_path: str, file_type: str = None) -> str:
        """
        解析文档为纯文本字符串
        保留来源元数据以便回溯
        """
        if file_type is None:
            file_type = self.get_file_type(file_path)

        if file_type not in self.SUPPORTED_TYPES:
            raise ValueError(f"不支持的文件类型: {file_type}")

        logger.info(f"开始解析文档: {file_path} (类型: {file_type})")

        try:
            if file_type == "pdf":
                text = self._parse_pdf(file_path)
            elif file_type == "word":
                text = self._parse_docx(file_path)
            elif file_type == "html":
                text = self._parse_html(file_path)
            elif file_type == "txt":
                text = self._parse_txt(file_path)

            logger.info(f"文档解析完成: {file_path}, 文本长度: {len(text)}")
            return text
        except Exception as e:
            logger.error(f"文档解析失败: {file_path}, 错误: {e}")
            raise

    def _parse_pdf(self, file_path: str) -> str:
        """使用 PyMuPDF 解析 PDF"""
        import fitz  # PyMuPDF

        text_parts = []
        doc = fitz.open(file_path)
        for page_num, page in enumerate(doc):
            text = page.get_text()
            if text.strip():
                text_parts.append(text.strip())
        doc.close()
        return "\n\n".join(text_parts)

    def _parse_docx(self, file_path: str) -> str:
        """使用 python-docx 解析 Word 文档"""
        from docx import Document

        doc = Document(file_path)
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())
        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)
        return "\n\n".join(text_parts)

    def _parse_html(self, file_path: str) -> str:
        """使用 BeautifulSoup4 解析 HTML"""
        from bs4 import BeautifulSoup

        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            html_content = f.read()
        soup = BeautifulSoup(html_content, "html.parser")
        # 移除 script 和 style 标签
        for tag in soup(["script", "style"]):
            tag.decompose()
        text = soup.get_text(separator="\n", strip=True)
        return text

    def _parse_txt(self, file_path: str) -> str:
        """直接读取文本文件"""
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read().strip()
